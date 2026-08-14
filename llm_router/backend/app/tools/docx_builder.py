"""Markdown → .docx builder（python-docx）。

供内置 ``generate_docx`` 工具使用：把智能体产出的 markdown 正文转成真正的 Word 二进制
（.docx），写入工作空间供下载，避免用 HTML 冒充 .docx。

支持 markdown 子集：``#``/``##``/``###`` 标题、段落（含 ``**加粗**`` 行内）、
``-``/``*`` 无序列表、``1.`` 有序列表、``| a | b |`` 管道表格。其余行按普通段落处理。
"""

from __future__ import annotations

import io
import re

from docx import Document

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\.?\.?\|?\s*$")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """把 ``**bold**`` 标记拆成多个 run；加粗段 ``bold=True``。"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _add_table(doc, table_lines: list[str]) -> None:
    # 丢弃 | --- | --- | 分隔行
    rows = [r for r in table_lines if not _TABLE_SEP_RE.match(r.strip())]
    if not rows:
        return
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    nrows = len(cells)
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(cells):
        for ci in range(ncols):
            table.cell(ri, ci).text = row[ci] if ci < len(row) else ""


def markdown_to_docx_bytes(markdown: str) -> bytes:
    """markdown 正文 → .docx 字节。"""
    doc = Document()
    lines = (markdown or "").splitlines()
    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        m = _HEADING_RE.match(stripped)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 3))
            i += 1
            continue
        bm = _BULLET_RE.match(stripped)
        if bm:
            doc.add_paragraph(bm.group(1), style="List Bullet")
            i += 1
            continue
        om = _ORDERED_RE.match(stripped)
        if om:
            doc.add_paragraph(om.group(1), style="List Number")
            i += 1
            continue
        if _is_table_row(stripped):
            table_lines: list[str] = []
            while i < n and _is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue
        para = doc.add_paragraph()
        _add_runs_with_bold(para, stripped)
        i += 1
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
